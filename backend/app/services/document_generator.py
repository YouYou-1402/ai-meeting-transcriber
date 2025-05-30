# backend/app/services/document_generator.py
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os
import logging
from datetime import datetime
from typing import Dict, Any, Optional, List


logger = logging.getLogger(__name__)

class DocumentGenerator:
    def __init__(self, output_folder: str):
        """
        Khởi tạo Document Generator
        
        Args:
            output_folder: Thư mục lưu file output
        """
        self.output_folder = output_folder
        os.makedirs(output_folder, exist_ok=True)
    
    def create_meeting_minutes(self, meeting_data: Dict[str, Any], filename: str = None) -> Optional[str]:
        """
        Tạo file Word biên bản cuộc họp
        
        Args:
            meeting_data: Dữ liệu cuộc họp
            filename: Tên file output (optional)
            
        Returns:
            Đường dẫn file đã tạo hoặc None nếu lỗi
        """
        try:
            # Tạo document mới
            doc = Document()
            
            # Cấu hình styles
            self._setup_document_styles(doc)
            
            # Thêm header
            self._add_header(doc, meeting_data)
            
            # Thêm thông tin chung
            self._add_general_info(doc, meeting_data)
            
            # Thêm nội dung chính
            self._add_main_content(doc, meeting_data)
            
            # Thêm quyết định
            self._add_decisions(doc, meeting_data)
            
            # Thêm nhiệm vụ cần làm
            self._add_action_items(doc, meeting_data)
            
            # Thêm phần kết
            self._add_footer_section(doc, meeting_data)
            
            # Tạo tên file nếu chưa có
            if not filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                title = meeting_data.get('title', 'meeting').replace(' ', '_')
                filename = f"bien_ban_{title}_{timestamp}.docx"
            
            # Đường dẫn file output
            output_path = os.path.join(self.output_folder, filename)
            
            # Lưu document
            doc.save(output_path)
            
            logger.info(f"Meeting minutes document created: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error creating meeting minutes document: {str(e)}")
            return None
    
    def _setup_document_styles(self, doc: Document):
        """Cấu hình styles cho document"""
        try:
            # Style cho tiêu đề chính
            title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = 'Times New Roman'
            title_font.size = Pt(16)
            title_font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(12)
            
            # Style cho tiêu đề phần
            heading_style = doc.styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
            heading_font = heading_style.font
            heading_font.name = 'Times New Roman'
            heading_font.size = Pt(12)
            heading_font.bold = True
            heading_style.paragraph_format.space_before = Pt(12)
            heading_style.paragraph_format.space_after = Pt(6)
            
            # Style cho nội dung
            content_style = doc.styles.add_style('CustomContent', WD_STYLE_TYPE.PARAGRAPH)
            content_font = content_style.font
            content_font.name = 'Times New Roman'
            content_font.size = Pt(11)
            content_style.paragraph_format.line_spacing = 1.15
            
        except Exception as e:
            logger.warning(f"Could not create custom styles: {str(e)}")
    
    def _add_header(self, doc: Document, meeting_data: Dict[str, Any]):
        """Thêm header cho document"""
        # Tiêu đề chính
        title = doc.add_paragraph()
        title_run = title.add_run("BIÊN BẢN CUỘC HỌP")
        title_run.font.name = 'Times New Roman'
        title_run.font.size = Pt(16)
        title_run.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Tên cuộc họp
        if meeting_data.get('title'):
            meeting_title = doc.add_paragraph()
            meeting_title_run = meeting_title.add_run(meeting_data['title'].upper())
            meeting_title_run.font.name = 'Times New Roman'
            meeting_title_run.font.size = Pt(14)
            meeting_title_run.bold = True
            meeting_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Thêm dòng trống
        doc.add_paragraph()
    
    def _add_general_info(self, doc: Document, meeting_data: Dict[str, Any]):
        """Thêm thông tin chung"""
        # Tiêu đề phần
        heading = doc.add_paragraph()
        heading_run = heading.add_run("I. THÔNG TIN CHUNG")
        heading_run.font.name = 'Times New Roman'
        heading_run.font.size = Pt(12)
        heading_run.bold = True
        
        # Thời gian
        time_info = meeting_data.get('created_at', datetime.now().strftime("%d/%m/%Y %H:%M"))
        if isinstance(time_info, str) and 'T' in time_info:
            # Convert ISO format to Vietnamese format
            try:
                dt = datetime.fromisoformat(time_info.replace('Z', '+00:00'))
                time_info = dt.strftime("%d/%m/%Y %H:%M")
            except:
                pass
        
        info_items = [
            f"- Thời gian: {time_info}",
            f"- Hình thức: Cuộc họp trực tuyến",
            f"- Thời lượng: {self._format_duration(meeting_data.get('duration', 0))}",
        ]
        
        # Người tham gia
        participants = meeting_data.get('participants', [])
        if participants:
            participants_str = ", ".join(participants)
            info_items.append(f"- Người tham gia: {participants_str}")
        else:
            info_items.append("- Người tham gia: [Danh sách người tham gia]")
        
        # Thêm thông tin vào document
        for item in info_items:
            p = doc.add_paragraph()
            run = p.add_run(item)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
        
        doc.add_paragraph()  # Dòng trống
    
    def _add_main_content(self, doc: Document, meeting_data: Dict[str, Any]):
        """Thêm nội dung chính của cuộc họp"""
        # Tiêu đề phần
        heading = doc.add_paragraph()
        heading_run = heading.add_run("II. NỘI DUNG CUỘC HỌP")
        heading_run.font.name = 'Times New Roman'
        heading_run.font.size = Pt(12)
        heading_run.bold = True
        
        # Tóm tắt nội dung
        summary = meeting_data.get('summary', '')
        if summary:
            # Tách summary thành các phần
            sections = self._parse_summary_sections(summary)
            
            for section_title, section_content in sections.items():
                if section_content.strip():
                    # Tiêu đề phần con
                    sub_heading = doc.add_paragraph()
                    sub_heading_run = sub_heading.add_run(f"2.{len(sections)}. {section_title}")
                    sub_heading_run.font.name = 'Times New Roman'
                    sub_heading_run.font.size = Pt(11)
                    sub_heading_run.bold = True
                    
                    # Nội dung
                    content_p = doc.add_paragraph()
                    content_run = content_p.add_run(section_content.strip())
                    content_run.font.name = 'Times New Roman'
                    content_run.font.size = Pt(11)
                    content_p.paragraph_format.left_indent = Inches(0.25)
        else:
            # Nếu không có summary, hiển thị transcript
            transcript = meeting_data.get('transcript', '')
            if transcript:
                content_p = doc.add_paragraph()
                content_run = content_p.add_run("Nội dung cuộc họp được ghi âm và chuyển đổi tự động:")
                content_run.font.name = 'Times New Roman'
                content_run.font.size = Pt(11)
                
                # Thêm transcript (giới hạn độ dài)
                transcript_preview = transcript[:2000] + "..." if len(transcript) > 2000 else transcript
                transcript_p = doc.add_paragraph()
                transcript_run = transcript_p.add_run(transcript_preview)
                transcript_run.font.name = 'Times New Roman'
                transcript_run.font.size = Pt(10)
                transcript_p.paragraph_format.left_indent = Inches(0.25)
        
        doc.add_paragraph()  # Dòng trống
    
    def _add_decisions(self, doc: Document, meeting_data: Dict[str, Any]):
        """Thêm các quyết định"""
        # Tiêu đề phần
        heading = doc.add_paragraph()
        heading_run = heading.add_run("III. CÁC QUYẾT ĐỊNH")
        heading_run.font.name = 'Times New Roman'
        heading_run.font.size = Pt(12)
        heading_run.bold = True
        
        # Trích xuất quyết định từ summary hoặc parsed data
        decisions = []
        
        # Thử lấy từ parsed data trước
        parsed_data = meeting_data.get('parsed_data', {})
        if parsed_data and 'decisions' in parsed_data:
            decisions = parsed_data['decisions']
        
        # Nếu không có, thử parse từ summary
        if not decisions:
            summary = meeting_data.get('summary', '')
            decisions = self._extract_decisions_from_summary(summary)
        
        if decisions:
            for i, decision in enumerate(decisions, 1):
                decision_p = doc.add_paragraph()
                decision_run = decision_p.add_run(f"{i}. {decision}")
                decision_run.font.name = 'Times New Roman'
                decision_run.font.size = Pt(11)
                decision_p.paragraph_format.left_indent = Inches(0.25)
        else:
            no_decisions_p = doc.add_paragraph()
            no_decisions_run = no_decisions_p.add_run("Không có quyết định cụ thể nào được đưa ra trong cuộc họp này.")
            no_decisions_run.font.name = 'Times New Roman'
            no_decisions_run.font.size = Pt(11)
            no_decisions_p.paragraph_format.left_indent = Inches(0.25)
        
        doc.add_paragraph()  # Dòng trống
    
    def _add_action_items(self, doc: Document, meeting_data: Dict[str, Any]):
        """Thêm nhiệm vụ cần làm"""
        # Tiêu đề phần
        heading = doc.add_paragraph()
        heading_run = heading.add_run("IV. NHIỆM VỤ CẦN LÀM")
        heading_run.font.name = 'Times New Roman'
        heading_run.font.size = Pt(12)
        heading_run.bold = True
        
        action_items = meeting_data.get('action_items', [])
        
        if action_items:
            # Tạo bảng cho action items
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            
            # Header row
            header_cells = table.rows[0].cells
            headers = ['STT', 'Nhiệm vụ', 'Người chịu trách nhiệm', 'Deadline']
            
            for i, header in enumerate(headers):
                header_cells[i].text = header
                # Format header
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11)
                        run.bold = True
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Data rows
            for i, item in enumerate(action_items, 1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(i)
                row_cells[1].text = item.get('task', '')
                row_cells[2].text = item.get('assignee', 'Chưa xác định')
                row_cells[3].text = item.get('deadline', 'Chưa xác định')
                
                # Format cells
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(10)
        else:
            no_actions_p = doc.add_paragraph()
            no_actions_run = no_actions_p.add_run("Không có nhiệm vụ cụ thể nào được giao trong cuộc họp này.")
            no_actions_run.font.name = 'Times New Roman'
            no_actions_run.font.size = Pt(11)
            no_actions_p.paragraph_format.left_indent = Inches(0.25)
        
        doc.add_paragraph()  # Dòng trống
    
    def _add_footer_section(self, doc: Document, meeting_data: Dict[str, Any]):
        """Thêm phần kết"""
        # Kết thúc cuộc họp
        heading = doc.add_paragraph()
        heading_run = heading.add_run("V. KẾT THÚC CUỘC HỌP")
        heading_run.font.name = 'Times New Roman'
        heading_run.font.size = Pt(12)
        heading_run.bold = True
        
        end_time = datetime.now().strftime("%H:%M ngày %d/%m/%Y")
        end_p = doc.add_paragraph()
        end_run = end_p.add_run(f"Cuộc họp kết thúc lúc {end_time}.")
        end_run.font.name = 'Times New Roman'
        end_run.font.size = Pt(11)
        end_p.paragraph_format.left_indent = Inches(0.25)
        
        # Chữ ký
        doc.add_paragraph()
        doc.add_paragraph()
        
        signature_table = doc.add_table(rows=3, cols=2)
        signature_table.autofit = False
        
        # Left column - Người ghi biên bản
        left_cells = [signature_table.cell(i, 0) for i in range(3)]
        left_cells[0].text = "NGƯỜI GHI BIÊN BẢN"
        left_cells[1].text = "(Ký tên)"
        left_cells[2].text = "[Tên người ghi]"
        
        # Right column - Chủ tọa cuộc họp
        right_cells = [signature_table.cell(i, 1) for i in range(3)]
        right_cells[0].text = "CHỦ TỌA CUỘC HỌP"
        right_cells[1].text = "(Ký tên)"
        right_cells[2].text = "[Tên chủ tọa]"
        
        # Format signature table
        for row in signature_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11)
                        if "NGƯỜI GHI" in run.text or "CHỦ TỌA" in run.text:
                            run.bold = True
    
    def _format_duration(self, duration_seconds: float) -> str:
        """Chuyển đổi thời lượng từ giây sang định dạng dễ đọc"""
        if not duration_seconds:
            return "Không xác định"
        
        hours = int(duration_seconds // 3600)
        minutes = int((duration_seconds % 3600) // 60)
        seconds = int(duration_seconds % 60)
        
        if hours > 0:
            return f"{hours} giờ {minutes} phút"
        elif minutes > 0:
            return f"{minutes} phút {seconds} giây"
        else:
            return f"{seconds} giây"
    
    def _parse_summary_sections(self, summary: str) -> Dict[str, str]:
        """Parse summary thành các sections"""
        sections = {}
        
        # Các pattern để tìm sections
        patterns = [
            (r'2\. MỤC ĐÍCH CUỘC HỌP:(.*?)3\. NỘI DUNG', 'Mục đích cuộc họp'),
            (r'3\. NỘI DUNG THẢO LUẬN:(.*?)4\. QUYẾT ĐỊNH', 'Nội dung thảo luận'),
            (r'6\. VẤN ĐỀ CẦN THEO DÕI:(.*?)7\. CUỘC HỌP', 'Vấn đề cần theo dõi'),
            (r'7\. CUỘC HỌP TIẾP THEO:(.*?)$', 'Cuộc họp tiếp theo')
        ]
        
        for pattern, section_name in patterns:
            import re
            match = re.search(pattern, summary, re.DOTALL)
            if match:
                sections[section_name] = match.group(1).strip()
        
        return sections
    
    def _extract_decisions_from_summary(self, summary: str) -> List[str]:
        """Trích xuất decisions từ summary"""
        decisions = []
        
        import re
        # Tìm phần quyết định
        decisions_match = re.search(r'4\. QUYẾT ĐỊNH:(.*?)5\. NHIỆM VỤ', summary, re.DOTALL)
        if decisions_match:
            decisions_text = decisions_match.group(1)
            # Tìm các items bắt đầu bằng -
            decision_items = re.findall(r'- (.*)', decisions_text)
            decisions = [item.strip() for item in decision_items if item.strip()]
        
        return decisions
    
    def create_transcript_document(self, meeting_data: Dict[str, Any], filename: str = None) -> Optional[str]:
        """
        Tạo file Word chứa transcript đầy đủ
        
        Args:
            meeting_data: Dữ liệu cuộc họp
            filename: Tên file output (optional)
            
        Returns:
            Đường dẫn file đã tạo hoặc None nếu lỗi
        """
        try:
            doc = Document()
            
            # Tiêu đề
            title = doc.add_paragraph()
            title_run = title.add_run("BẢN GHI ÂM CUỘC HỌP")
            title_run.font.name = 'Times New Roman'
            title_run.font.size = Pt(16)
            title_run.bold = True
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Thông tin cuộc họp
            if meeting_data.get('title'):
                meeting_title = doc.add_paragraph()
                meeting_title_run = meeting_title.add_run(meeting_data['title'])
                meeting_title_run.font.name = 'Times New Roman'
                meeting_title_run.font.size = Pt(14)
                meeting_title_run.bold = True
                meeting_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()
            
            # Thông tin file
            info_p = doc.add_paragraph()
            info_run = info_p.add_run(f"File: {meeting_data.get('filename', 'N/A')}")
            info_run.font.name = 'Times New Roman'
            info_run.font.size = Pt(11)
            
            duration_p = doc.add_paragraph()
            duration_run = duration_p.add_run(f"Thời lượng: {self._format_duration(meeting_data.get('duration', 0))}")
            duration_run.font.name = 'Times New Roman'
            duration_run.font.size = Pt(11)
            
            doc.add_paragraph()
            
            # Transcript content
            transcript = meeting_data.get('transcript', '')
            if transcript:
                transcript_heading = doc.add_paragraph()
                transcript_heading_run = transcript_heading.add_run("NỘI DUNG BẢN GHI:")
                transcript_heading_run.font.name = 'Times New Roman'
                transcript_heading_run.font.size = Pt(12)
                transcript_heading_run.bold = True
                
                # Thêm transcript
                transcript_p = doc.add_paragraph()
                transcript_run = transcript_p.add_run(transcript)
                transcript_run.font.name = 'Times New Roman'
                transcript_run.font.size = Pt(11)
                transcript_p.paragraph_format.line_spacing = 1.5
            
            # Tạo tên file
            if not filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                title = meeting_data.get('title', 'transcript').replace(' ', '_')
                filename = f"transcript_{title}_{timestamp}.docx"
            
            output_path = os.path.join(self.output_folder, filename)
            doc.save(output_path)
            
            logger.info(f"Transcript document created: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error creating transcript document: {str(e)}")
            return None
